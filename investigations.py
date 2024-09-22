from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

# Function to set cell borders
def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for border_name, border_attrs in kwargs.items():
        el = OxmlElement('w:' + border_name)
        for key, value in border_attrs.items():
            el.set(qn('w:' + key), str(value))
        tcPr.append(el)

# Function to set table width
def set_table_width(table, width):
    tbl = table._element
    tblPr = tbl.xpath("w:tblPr")
    if not tblPr:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    else:
        tblPr = tblPr[0]
    
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

# Function to create the patient document
def create_patient_document(investigations, advise):
    doc = Document()
    
    # Create the main table with enough columns for all dates plus one for the investigation and one for advise
    num_dates = len(investigations['dates'])
    main_table = doc.add_table(rows=1, cols=num_dates + 2)
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set the total width of the table to 19.57 cm
    total_width_cm = 19.57
    total_width_twips = int(total_width_cm * 567)  # Convert cm to twips

    # Set table width and column widths
    set_table_width(main_table, total_width_twips)
    inv_width_twips = int(total_width_twips * 0.65)  # 65% for investigations
    adv_width_twips = total_width_twips - inv_width_twips
    main_table.columns[0].width = Cm(inv_width_twips / 567)
    for i in range(1, num_dates + 1):
        main_table.columns[i].width = Cm((inv_width_twips / 567) / num_dates)
    main_table.columns[num_dates + 1].width = Cm(adv_width_twips / 567)

    # Add headers
    hdr_cells = main_table.rows[0].cells
    hdr_cells[0].text = 'Investigation'
    for i, date in enumerate(investigations['dates']):
        hdr_cells[i + 1].text = date
    hdr_cells[num_dates + 1].text = 'Advise on Discharge'

    # Bold the header cells
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(9)

    # Fill investigation data
    for test_name, results in investigations['data'].items():
        row_cells = main_table.add_row().cells
        row_cells[0].text = test_name
        for run in row_cells[0].paragraphs[0].runs:
            run.bold = True  # Bold the test name tags
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
        for i, result in enumerate(results):
            row_cells[i + 1].text = result

    # Set borders for each cell in the last column before merging
    for row in main_table.rows[1:]:
        set_cell_border(row.cells[-1], top={"sz": "12", "val": "single"}, bottom={"sz": "12", "val": "single"},
                        start={"sz": "12", "val": "single"}, end={"sz": "12", "val": "single"})

    # Merge all cells in the last column for advise
    for row in main_table.rows[1:]:  # Skip header row
        row.cells[-1].merge(main_table.cell(1, num_dates + 1))

    # Fill advise section in the merged cell with line spacing
    merged_adv_cell = main_table.cell(1, num_dates + 1)
    paragraphs = merged_adv_cell.paragraphs
    if len(paragraphs) == 0:
        merged_adv_cell.add_paragraph()  # Ensure there's at least one paragraph to start

    # Adding small headings for each section
    sections = {
    'General': advise.get('general', None),
    'Medications': advise.get('medications', None),
    'Follow Up': advise.get('follow_up', None)
    }


    for section_name, items in sections.items():
        # Add heading for the section
        para = merged_adv_cell.add_paragraph(section_name)
        para.paragraph_format.space_after = Pt(6)  # Small spacing after the heading
        for run in para.runs:
            run.bold = True  # Bold the section headings
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)

        # Add items as bullet points
        for item in items:
            para = merged_adv_cell.add_paragraph(item, style='ListBullet')
            para.paragraph_format.space_after = Pt(2)  # Small spacing after each bullet point
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(9)

    # Add the doctor's name at the end of the advise section
    doctor_para = merged_adv_cell.add_paragraph("Doctor: Taruna Pahuja (Nephrology SR)")
    doctor_para.paragraph_format.space_after = Pt(10)  # Add spacing after the doctor's name
    for run in doctor_para.runs:
        run.bold = True  # Bold the doctor's name
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)

    # Add space below the doctor's name for signature
    signature_para = merged_adv_cell.add_paragraph()
    signature_para.paragraph_format.space_after = Pt(30)  # Add extra space for signature
    for run in signature_para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)

    # Reapply borders after merging
    set_cell_border(merged_adv_cell, top={"sz": "12", "val": "single"}, bottom={"sz": "12", "val": "single"},
                    start={"sz": "12", "val": "single"}, end={"sz": "12", "val": "single"})

    # Apply font, border style, and bold header for all cells
    for row in main_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)
            set_cell_border(cell, top={"sz": "12", "val": "single"}, bottom={"sz": "12", "val": "single"},
                            start={"sz": "12", "val": "single"}, end={"sz": "12", "val": "single"})
    
    # Save the document
    doc.save('output_dynamic_streamlit.docx')
    doc_filled = Document("template_patient_info.docx")
    doc_dynamic = Document("output_dynamic_streamlit.docx")

# Append the content of output_dynamic to output_filled
    for element in doc_dynamic.element.body:
        doc_filled.element.body.append(element)

# Save the updated document
    doc_filled.save("final_dis_summary.docx")
